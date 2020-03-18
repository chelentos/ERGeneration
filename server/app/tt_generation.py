import datetime

from docx import Document
from docx.shared import Inches

from flask_login import current_user

from pathlib import Path

reqsTypes = {
  "F": "Функциональные требования",
  "FT": "Требования к отказоустойчивости",
  "LF": "UI",
  "SE": "Требования к безопасности",
  "SC": "Требования к масштабируемости",
  "A": "Требования к доступности",
  "O": "Требования к поддерживаемости",
  "US": "UX",
  "L": "Юридические требования",
  "PE": "Требования к производительности"
}

def generateTT(reqs):
  document = Document()

  document.add_heading('Техническое задание', 0)

  for t in reqsTypes.keys():
    print(t)
    typedReqs = list(filter(lambda r: r['type'] == t, reqs))
    print(typedReqs)
    if len(typedReqs) > 0:
      document.add_heading(reqsTypes[t], 1)
      for tr in typedReqs:
        document.add_paragraph(
          tr['text'], style='List Bullet'
        )
  now = datetime.datetime.now()
  Path("app/public/" + str(current_user.id)).mkdir(parents=True, exist_ok=True)
  document.save('app/public/' + str(current_user.id) + '/tt_' + now.strftime("%Y.%m.%d_%H.%M.%S") + '.docx')
  return 'public/' + str(current_user.id) + '/tt_' + now.strftime("%Y.%m.%d_%H.%M.%S") + '.docx'



